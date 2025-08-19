from flask import Flask, request
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import requests
import tempfile
import subprocess
import os
import sys
import time
import win32com.client
import pythoncom




app = Flask(__name__)
def create_zebra_label(c_data):
    return f"""
^XA
^PW812
^LL1218
^MD30

^FO40,50^A0N,60,60^FDName: {c_data.get('customer_name','N/A')}^FS
^FO40,125^A0N,60,60^FDPhone: {c_data.get('phone_number','N/A')}^FS
^FO40,200^A0N,60,60^FDYear: {c_data.get('year','N/A')}^FS
^FO40,275^A0N,60,60^FDMake: {c_data.get('make','N/A')}^FS
^FO40,350^A0N,60,60^FDModel: {c_data.get('model','N/A')}^FS
^FO40,425^A0N,60,60^FDDiag: {c_data.get('diagnostics','N/A')}^FS
^FO40,500^A0N,60,60^FDModules: {c_data.get('modules','N/A')}^FS
^FO40,575^A0N,60,60^FDSingle: {c_data.get('single_stage','N/A')}^FS
^FO40,650^A0N,60,60^FDDual: {c_data.get('dual_stage','N/A')}^FS
^FO40,725^A0N,60,60^FDBuckles: {c_data.get('buckles','N/A')}^FS
^FO40,800^A0N,60,60^FDDrop-off: {c_data.get('wwydoff','N/A')}^FS


^FO260,960
^BQN,2,4
^FDMA,https://forms.gle/C9z4jsXJas8gR5oa9^FS

^XZ



""".strip()


# def print_and_close_word(doc_path):
#     print("Opening word...")
#     pythoncom.CoInitialize()# needs to be initialize manually since ngrok is
#     #running on a thread, word doesnt like to initialize from a thread so
#     # one has to initialize manually
    
#     word = win32com.client.Dispatch("Word.Application")
#     word.Visible = True
    
#     try:
#         doc = word.Documents.Open(doc_path)
#         doc.PrintOut(Background=False, Range=0)

#         time.sleep(5)
#         doc.Close(False)
#         print("☑️Document printed and closed")
#     except Exception as e:
#         print(f"❌ Error while printing:{e}")
#     finally:
#         word.Quit()
#         print("📁 word program closed.")
# def add_bolded_content(doc, content):
#     for line in content.strip().split('\n'):
#         if ':' in line:
#             label, value = line.split(':', 1)
#             paragraph = doc.add_paragraph()
#             run1 = paragraph.add_run(label + ": ")
#             run1.font.size = Pt(14)

#             run2 = paragraph.add_run(value.strip())
#             run2.bold = True
#             run2.font.size = Pt(14)
#         else:
#             paragraph = doc.add_paragraph(line)

# ✅ Printer function using Notepad to print the content
# def print_to_printer(content):
#     doc = Document()
#     paragraph = doc.add_paragraph()
#     run= paragraph.add_run(content)
#     run.font.size= Pt(14)
  


#     if getattr(sys, 'frozen', False):
#         script_dir = sys._MEIPASS
#     else:
#         script_dir = os.path.dirname(os.path.abspath(__file__))
#     image_path = os.path.join(script_dir, "app", "Sf_QR.JPEG")
#     doc.add_picture(image_path,width=Inches(1.00),height=Inches(1.00))
#     # print("the paragraph:",paragraph.text)
#     # Write the content to a temporary text file replace to .docx
#     with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmpfile:
       
#         tmpfile_path = tmpfile.name

#     print(f"📄 Temp file created at: {tmpfile_path}")
#     doc.save(tmpfile_path)
#     word_path =r"C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE"
#     print_and_close_word(tmpfile_path)
    # try:
    #     print_and_close_word(tempfile_path)
    #     # Use Notepad to print the file
    # #     subprocess.run([word_path,"/q","/n","/mFilePrintDefault","/t",tmpfile_path], check = True)
    # #     print("✅ Print job sent via notepad")
    # # except subprocess.CalledProcessError as e:
    # #     print(f"❌ Failed to print:, {e.returncode}")
    
    # except Exception as e:
    #     print(f"❌ unexpected error {e}")

        

    # Optional: delete the temp file after printing
    # time.sleep(5)  # Give word time to send the print job
    # try:
    #     os.remove(tmpfile_path)
    #     print("🗑️ Temp file deleted.")
    # except Exception as e:
    #     print("⚠️ Could not delete temp file:", e)
def zend_zpl_through_usb(zpl):
    # turning text commands into raw bytes usb understands zpl.encode("utf-8"), change w to "wb"
    port = r"\\.\USB001"  # change to your actual one (USB001/USB003/etc.)
    try:
        with open(port, "wb") as printer:
            printer.write(zpl.encode("utf-8"))
    except FileNotFoundError:
        raise RuntimeError(f"Port {port} not found. Check Printer Properties → Ports.")
    except PermissionError:
        raise RuntimeError(f"Permission denied opening {port}. Try running as admin.")



# ✅ /test endpoint to verify connection
@app.route('/test', methods=['GET'])
def test_connection():
    return {'status': '✅ Flask is working!'}

# ✅ /print endpoint to receive data and trigger printing
@app.route('/print', methods=['POST'])
def print_info():
   
    try:
        data = request.get_json(force=True)
        print("📥 Received data:", data)
        zpl=create_zebra_label(data)
        zend_zpl_through_usb(zpl)
        return f" printer create zebra data{data}"

        # ✅ Rebuild the message from timestamp, name, phone, service
#         content = f"""🕒 Time: {data['timestamp']}
# 👤 Name: {data['customer_name']}
# 📞 Phone#: {data['phone_number']}
#     Year: {data['year']}
#     Make: {data['make']}
#     Model: {data['model']}
#     Diagnostics: {data['diagnostics']}
#     Modules: {data['modules']}
#     Single_stage: {data['single_stage']}
#     Dual_stage: {data['dual_stage']}
#     Buckles: {data['buckles']}
#     what_will_you_drop_off: {data['wwydoff']}
    
# """

#         print("📝 Content to print:", content)

#         # print_to_printer(content)
     
#         return {'status': 'printed'}
    except Exception as e:
        print("❌ Error in /print route:", e)
        return {'status': 'error', 'message': str(e)}, 500

       


def post_url_to_google_apps_script(ngrok_url):
    webhook_url = "https://script.google.com/macros/s/AKfycbxOx28J1NCWW_cN9Vk81Cb4xYcb49aNkk1FnprS9HCFCd9dhRH3w_a24Ibv4Df2_rlV/exec"
    requests.post(webhook_url, json={"url": ngrok_url})

# ✅ Run Flask on 0.0.0.0 so it works with ngrok
if __name__ == "__main__":
    # url = start_ngrok_server()
    # if url:
    #     post_url_to_google_apps_script(url)
    # app.run(host='0.0.0.0', port=5000, threaded=True, debug=True)
    app.run(host='0.0.0.0', port=5000, debug=True)





